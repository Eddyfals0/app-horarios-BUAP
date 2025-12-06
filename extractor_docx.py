import csv
import os
import re
from docx import Document

def clean_text(text):
    return text.strip().replace('\n', ' ').replace('\r', '')

def extract_metadata_from_text(header_text):
    carrera = "Desconocida"
    recinto = "Desconocido"
    
    # Extract Carrera
    match_carrera = re.search(r'((?:LICENCIATURA|INGENIERÍA|MAESTRÍA|DOCTORADO) EN [A-ZÁÉÍÓÚÑ ]+?)(?=\s+PROGRAMACIÓN| [A-ZÁÉÍÓÚÑ]+$)', header_text)
    if match_carrera:
        carrera = match_carrera.group(1).strip()
    else:
        # Fallback split
        for keyword in ["LICENCIATURA", "INGENIERÍA", "MAESTRÍA", "DOCTORADO"]:
            if keyword in header_text:
                 parts = header_text.split(keyword)
                 if len(parts) > 1:
                     sub = keyword + parts[1]
                     carrera = sub.split("PROGRAMACIÓN")[0].strip()
                     break

    # Extract Recinto
    header_upper = header_text.upper()
    if "CU SAN MANUEL" in header_upper or "CIUDAD UNIVERSITARIA" in header_upper:
        recinto = "CU"
    elif "CU2" in header_upper or "ECOCAMPUS" in header_upper:
        recinto = "CU2"
    else:
        if "CU2" in header_upper:
            recinto = "CU2"
        elif "CU" in header_upper:
            recinto = "CU"
            
    return carrera.title(), recinto

def is_header_row(row_cells):
    # Detect if a row is the column headers row
    return "NRC" in row_cells and "Materia" in row_cells

def is_metadata_row(row_cells):
    # Detect if a row is the top-level metadata row (e.g. FACULTAD DE CIENCIAS...)
    # usually the first cell is very long and contains "FACULTAD"
    text = " ".join(row_cells).upper()
    return "FACULTAD" in text and "PROGRAMACIÓN ACADÉMICA" in text

def docx_to_csv(input_path):
    if not os.path.exists(input_path):
        print(f"Error: File not found at {input_path}")
        return

    try:
        doc = Document(input_path)
    except Exception as e:
        print(f"Error opening docx: {e}")
        return

    all_data = []
    headers = []
    
    global_carrera = "Desconocida"
    global_recinto = "Desconocido"
    metadata_found = False

    print(f"Total tables found: {len(doc.tables)}")

    for t_idx, table in enumerate(doc.tables):
        rows = table.rows
        for r_idx, row in enumerate(rows):
            cells = [clean_text(cell.text) for cell in row.cells]
            
            # Skip empty rows
            if not any(cells):
                continue

            # Check for Metadata
            if is_metadata_row(cells):
                # If we haven't found metadata yet, extract it
                if not metadata_found:
                    full_text = cells[0] 
                    # Sometimes the text might be split or in a different cell
                    if not full_text: full_text = " ".join(cells)
                    
                    global_carrera, global_recinto = extract_metadata_from_text(full_text)
                    print(f"Metadata identified: {global_carrera} ({global_recinto})")
                    metadata_found = True
                continue # Skip metadata row in output

            # Check for Headers
            if is_header_row(cells):
                if not headers:
                    headers = cells
                    print(f"Headers identified: {headers}")
                continue # Skip header row in data
            
            # Assume it's data
            all_data.append(cells)

    if not all_data:
        print("No data extracted.")
        return

    # Generate Output Filename
    safe_carrera = re.sub(r'[<>:"/\\|?*]', '', global_carrera)
    safe_recinto = re.sub(r'[<>:"/\\|?*]', '', global_recinto)
    
    output_filename = f"{safe_carrera} - {safe_recinto}.csv"
    output_dir = os.path.dirname(input_path)
    output_path = os.path.join(output_dir, output_filename)

    # Write CSV
    try:
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            if headers:
                writer.writerow(headers)
            else:
                # Fallback headers if none found
                writer.writerow(['NRC', 'Clave', 'Materia', 'Secc', 'Dias', 'Hora', 'Profesor', 'Salon'])
            
            writer.writerows(all_data)
        print(f"Successfully created: {output_path}")
        print(f"Total rows extracted: {len(all_data)}")
    except Exception as e:
        print(f"Error writing CSV: {e}")

if __name__ == "__main__":
    path = r"C:\Users\Eduar\Documents\GitHub\app-horarios-BUAP\Pruba de extración de información\PA_ICC_CU_21_NOV_2025.docx"
    docx_to_csv(path)
