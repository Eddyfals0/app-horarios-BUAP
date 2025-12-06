from docx import Document
import os

def inspect_header(file_path):
    print(f"Inspecting: {file_path}")
    try:
        doc = Document(file_path)
        if doc.tables:
            table = doc.tables[0]
            if table.rows:
                # Print first 3 rows
                for i, row in enumerate(table.rows[:3]):
                    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    print(f"Row {i}: {cells}")
        else:
            print("No tables found.")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    path = r"C:\Users\Eduar\Documents\GitHub\app-horarios-BUAP\Pruba de extración de información\PA_ICC_CU_21_NOV_2025.docx"
    inspect_header(path)
